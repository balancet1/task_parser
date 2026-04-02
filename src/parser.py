import pdfplumber
import re
from datetime import datetime
from typing import List, Dict, Optional
import os
import subprocess
import shutil

class TaskParser:
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.tasks = []
        self.file_extension = os.path.splitext(file_path)[1].lower()
        
        # ========== НАСТРОЙКИ ПОЛЕЙ (МОЖНО МЕНЯТЬ) ==========
        # Ключевые слова для поиска даты
        self.date_keywords = ['Срок', 'Дата', 'Дедлайн', 'Due', 'Выполнить до']
        
        # Ключевые слова для поиска ответственного
        self.resp_keywords = ['Отв.', 'Исполнитель', 'Ответственный', 'Исп.']
        
        # Текстовые статусы выполнения
        self.status_keywords = ['выполнено', 'выполнен', 'сделано', 'готово']
        
        # Разделители между словом и значением
        self.separators = r'\s*(?:—|–|-|:)?\s*'
        
        # Слова, которые означают конец раздела с задачами
        self.after_keywords = [
            'Протокол вёл', 
            'Лист согласования', 
            'Стр.', 
            'Page',
            'Ознакомлены',
            'Подписи',
            'УТВЕРЖДАЮ',
            'СОГЛАСОВАНО',
            'От АО «ТАНЕКО»:',
            'От ООО «НТЦ Татнефть»:',
            'От ООО «ЭПИК»:'
        ]
        
        # Слова, которые игнорируются до РЕШИЛИ:
        self.before_keywords = [
            'СЛУШАЛИ:', 
            'ВЫСТУПИЛИ:', 
            'ДОКЛАДЫВАЛИ:', 
            'ОБСУЖДАЛИ:',
            'ПОВЕСТКА ДНЯ:',
            'ПРИСУТСТВОВАЛИ:'
        ]
        # ====================================================
    
    def extract_text(self) -> str:
        """Извлекает текст из файла (поддерживает PDF, DOCX и DOC)"""
        
        if self.file_extension == '.pdf':
            return self._extract_from_pdf()
        elif self.file_extension == '.docx':
            return self._extract_from_docx()
        elif self.file_extension == '.doc':
            return self._extract_from_doc()
        else:
            print(f"❌ Неподдерживаемый формат файла: {self.file_extension}")
            print("   Поддерживаются: .pdf, .docx, .doc")
            return ""
    
    def _extract_from_pdf(self) -> str:
        full_text = ""
        try:
            with pdfplumber.open(self.file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        full_text += text + "\n"
            print(f"✅ Извлечено {len(full_text)} символов из PDF")
            return full_text
        except Exception as e:
            print(f"❌ Ошибка при чтении PDF: {e}")
            return ""
    
    def _extract_from_docx(self) -> str:
        try:
            from docx import Document
            
            doc = Document(self.file_path)
            full_text = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    try:
                        import xml.etree.ElementTree as ET
                        if para._element.xpath('.//w:numPr'):
                            full_text.append(f"¶ {text}")
                        else:
                            full_text.append(text)
                    except:
                        full_text.append(text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        full_text.append(' | '.join(row_text))
            
            result = '\n'.join(full_text)
            print(f"✅ Извлечено {len(result)} символов из Word документа (.docx)")
            return result
            
        except ImportError:
            print("❌ Библиотека python-docx не установлена")
            print("   Установите: pip install python-docx")
            return ""
        except Exception as e:
            print(f"❌ Ошибка при чтении Word документа: {e}")
            return ""
    
    def _extract_from_doc(self) -> str:
        if shutil.which('antiword'):
            try:
                result = subprocess.run(['antiword', self.file_path], 
                                       capture_output=True, text=True)
                if result.returncode == 0:
                    print(f"✅ Извлечено {len(result.stdout)} символов из Word .doc файла")
                    return result.stdout
            except Exception as e:
                print(f"⚠️ Ошибка antiword: {e}")
        
        if shutil.which('soffice'):
            try:
                import tempfile
                temp_dir = tempfile.mkdtemp()
                
                result = subprocess.run([
                    'soffice', '--headless', '--convert-to', 'txt',
                    '--outdir', temp_dir, self.file_path
                ], capture_output=True, text=True)
                
                if result.returncode == 0:
                    base_name = os.path.basename(self.file_path).replace('.doc', '.txt')
                    txt_file = os.path.join(temp_dir, base_name)
                    
                    if os.path.exists(txt_file):
                        with open(txt_file, 'r', encoding='utf-8', errors='ignore') as f:
                            content = f.read()
                        
                        os.remove(txt_file)
                        os.rmdir(temp_dir)
                        print(f"✅ Извлечено {len(content)} символов из Word .doc файла (через LibreOffice)")
                        return content
            except Exception as e:
                print(f"⚠️ Ошибка при конвертации через LibreOffice: {e}")
        
        print("❌ Не удалось извлечь текст из .doc файла.")
        print("   Установите: brew install antiword")
        return ""
    
    def parse_tasks(self, text: str) -> List[Dict]:
        lines = text.split('\n')
        
        has_resheno = False
        resheno_index = -1
        for i, line in enumerate(lines[:100]):
            if 'РЕШИЛИ:' in line:
                has_resheno = True
                resheno_index = i
                print(f"✅ Найден маркер 'РЕШИЛИ:' в строке {i}")
                break
        
        if self.file_extension == '.pdf':
            print("📄 PDF файл: использую простой парсинг")
            self.tasks = self._parse_pdf_simple(lines)
        elif has_resheno:
            print("📝 Word файл с РЕШИЛИ: использую парсинг протокола")
            self.tasks = self._parse_word_protocol(lines, resheno_index)
        else:
            print("📄 Простой список: использую базовый парсинг")
            self.tasks = self._parse_simple_list(lines)
        
        return self.tasks
    
    def _parse_pdf_simple(self, lines: List[str]) -> List[Dict]:
        tasks = []
        current_task = None
        current_description = []
        
        решили_index = -1
        for i, line in enumerate(lines):
            if 'РЕШИЛИ:' in line:
                решили_index = i
                break
        
        start_index = решили_index + 1 if решили_index != -1 else 0
        
        i = start_index
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue
            
            stop_parsing = False
            for keyword in self.after_keywords:
                if keyword in line[:30]:
                    stop_parsing = True
                    break
            if stop_parsing:
                break
            
            task_match = re.match(r'^(\d+)\.\s+(.*)', line)
            
            if task_match:
                if current_task:
                    full_desc = ' '.join(current_description)
                    full_desc = re.sub(r'\s+', ' ', full_desc)
                    
                    # Очищаем описание от метаданных
                    full_desc = re.sub(r'[;,\s]*Отв\.:\s*[^\.]+?(?:\.|$)', '', full_desc)
                    full_desc = re.sub(r'[;,\s]*Отв\.:\s*[^С]+?(?:\s+Срок|$)', '', full_desc)
                    full_desc = re.sub(r'[;,\s]*Отв\.:\s*[^\n]+', '', full_desc)
                    full_desc = re.sub(r'[;,\s]*Срок\s*[—–-]?\s*\d{2}\.\d{2}\.\d{4}', '', full_desc)
                    full_desc = re.sub(r'[;,\s]*Срок\s*[—–-]?\s*до\s+конца\s+года', '', full_desc)
                    full_desc = re.sub(r'[;,\s]*С\b', '', full_desc)
                    full_desc = re.sub(r'\s+', ' ', full_desc)
                    full_desc = re.sub(r'\s*[;,]?\s*$', '', full_desc)
                    full_desc = full_desc.strip()
                    
                    current_task['full_description'] = full_desc
                    tasks.append(current_task)
                
                task_num = task_match.group(1)
                task_text = task_match.group(2)
                
                current_task = {
                    'number': int(task_num),
                    'full_description': '',
                    'responsible': '',
                    'due_date': None,
                    'due_date_str': ''
                }
                current_description = [task_text]
                i += 1
            
            elif current_task:
                current_description.append(line)
                
                if 'Отв.:' in line:
                    resp_match = re.search(r'Отв\.:\s*([^С]+?)(?:\s+Срок|$)', line)
                    if not resp_match:
                        resp_match = re.search(r'Отв\.:\s*([^\n]+)', line)
                    
                    if resp_match:
                        responsible = resp_match.group(1).strip()
                        
                        # Обрезаем до ключевых слов
                        stop_words = self.date_keywords + ['Выполнено', 'Приложение', 'приложение', 'Протокол'] + self.status_keywords
                        
                        for stop_word in stop_words:
                            if stop_word in responsible:
                                responsible = responsible.split(stop_word)[0].strip()
                                break
                        
                        # Дополнительная очистка от "Срок" и "Выполнено" в любом регистре
                        responsible = re.sub(r'\s+Срок.*$', '', responsible, flags=re.IGNORECASE)
                        responsible = re.sub(r'\s+Выполнено.*$', '', responsible, flags=re.IGNORECASE)
                        responsible = re.sub(r'\s+до\s+конца\s+года.*$', '', responsible, flags=re.IGNORECASE)
                        
                        responsible = re.sub(r'\s+', ' ', responsible)
                        current_task['responsible'] = responsible
                
                if 'Срок' in line or any(word in line.lower() for word in self.status_keywords + ['до конца года']):
                    line_lower = line.lower()
                    
                    # Проверяем на статусы выполнения
                    if any(word in line_lower for word in self.status_keywords):
                        current_task['due_date_str'] = 'Выполнено'
                    elif 'до конца года' in line_lower:
                        current_task['due_date_str'] = 'до конца года'
                    else:
                        date_match = re.search(r'Срок\s*[—–-]?\s*(\d{2}\.\d{2}\.\d{4})', line)
                        if date_match:
                            date_str = date_match.group(1).strip()
                            current_task['due_date_str'] = date_str
                            try:
                                current_task['due_date'] = datetime.strptime(date_str, '%d.%m.%Y').date()
                            except ValueError:
                                pass
                
                i += 1
            else:
                i += 1
    
        if current_task:
            full_desc = ' '.join(current_description)
            full_desc = re.sub(r'\s+', ' ', full_desc)
            
            # Очищаем описание от метаданных
            full_desc = re.sub(r'[;,\s]*Отв\.:\s*[^\.]+?(?:\.|$)', '', full_desc)
            full_desc = re.sub(r'[;,\s]*Отв\.:\s*[^С]+?(?:\s+Срок|$)', '', full_desc)
            full_desc = re.sub(r'[;,\s]*Отв\.:\s*[^\n]+', '', full_desc)
            full_desc = re.sub(r'[;,\s]*Срок\s*[—–-]?\s*\d{2}\.\d{2}\.\d{4}', '', full_desc)
            full_desc = re.sub(r'[;,\s]*Срок\s*[—–-]?\s*до\s+конца\s+года', '', full_desc)
            full_desc = re.sub(r'[;,\s]*С\b', '', full_desc)
            full_desc = re.sub(r'\s+', ' ', full_desc)
            full_desc = re.sub(r'\s*[;,]?\s*$', '', full_desc)
            full_desc = full_desc.strip()
            
            current_task['full_description'] = full_desc
            tasks.append(current_task)
        
        return tasks
    
    def _parse_word_protocol(self, all_lines: List[str], start_idx: int) -> List[Dict]:
        tasks = []
        
        решили_pos = -1
        for i, line in enumerate(all_lines):
            if 'РЕШИЛИ:' in line:
                решили_pos = i
                break
        
        if решили_pos == -1:
            return []
        
        task_lines = []
        i = решили_pos + 1
        
        while i < len(all_lines) and not all_lines[i].strip():
            i += 1
        
        started = False
        
        while i < len(all_lines):
            line = all_lines[i].strip()
            
            stop_found = False
            for keyword in self.after_keywords:
                if keyword in line[:30]:
                    stop_found = True
                    break
            
            if stop_found:
                break
            
            if re.match(r'^\d+$', line):
                i += 1
                continue
            
            is_service = False
            for keyword in self.before_keywords:
                if keyword in line:
                    is_service = True
                    break
            
            if is_service:
                i += 1
                continue
            
            if not started and (re.match(r'^\d+[.\t]', line) or line.startswith('¶')):
                started = True
            
            if started and line:
                task_lines.append(line)
            
            i += 1
        
        if task_lines:
            last_line = task_lines[-1]
            for keyword in self.after_keywords:
                if keyword in last_line:
                    task_lines[-1] = last_line.split(keyword)[0].strip()
                    break
        
        if task_lines and not task_lines[-1]:
            task_lines.pop()
        
        i = 0
        task_counter = 1
        
        while i < len(task_lines):
            line = task_lines[i]
            
            is_task_start = False
            task_number = None
            description = None
            
            match = re.match(r'^(\d+)[.\t]\s*(.*)', line)
            if match:
                is_task_start = True
                task_number = int(match.group(1))
                description = match.group(2)
            
            if not is_task_start and line.startswith('¶'):
                is_task_start = True
                task_number = task_counter
                description = re.sub(r'^¶\s*', '', line)
            
            if not is_task_start:
                has_resp = any(k in line for k in self.resp_keywords)
                has_date = any(k in line for k in self.date_keywords)
                
                is_service = False
                for keyword in self.before_keywords:
                    if keyword in line:
                        is_service = True
                        break
                
                if not has_resp and not has_date and not is_service and len(line) > 20:
                    is_task_start = True
                    task_number = task_counter
                    description = line
            
            if is_task_start and description:
                i += 1
                responsible = ""
                due_date_str = ""
                due_date = None
                
                while i < len(task_lines) and not task_lines[i].strip():
                    i += 1
                
                collected_resp = False
                collected_date = False
                
                while i < len(task_lines) and not (collected_resp and collected_date):
                    current = task_lines[i].strip()
                    
                    if not current:
                        i += 1
                        continue
                    
                    next_is_task = False
                    if re.match(r'^\d+[.\t]', current):
                        next_is_task = True
                    elif current.startswith('¶'):
                        next_is_task = True
                    else:
                        has_resp_next = any(k in current for k in self.resp_keywords)
                        has_date_next = any(k in current for k in self.date_keywords)
                        is_service_next = any(k in current for k in self.before_keywords)
                        if not has_resp_next and not has_date_next and not is_service_next and len(current) > 20:
                            next_is_task = True
                    
                    if next_is_task:
                        break
                    
                    if not collected_resp:
                        for keyword in self.resp_keywords:
                            if keyword in current:
                                resp_parts = current.split(keyword)
                                if len(resp_parts) > 1:
                                    resp_text = resp_parts[1].strip()
                                    
                                    stop_words = self.date_keywords + ['Выполнено', 'Приложение', 'приложение', 'Протокол'] + self.status_keywords
                                    
                                    for stop_word in stop_words:
                                        if stop_word in resp_text.lower():
                                            resp_text = resp_text.split(stop_word)[0].strip()
                                            break
                                    
                                    # Дополнительная очистка
                                    resp_text = re.sub(r'\s+Срок.*$', '', resp_text, flags=re.IGNORECASE)
                                    resp_text = re.sub(r'\s+Выполнено.*$', '', resp_text, flags=re.IGNORECASE)
                                    resp_text = re.sub(r'\s+до\s+конца\s+года.*$', '', resp_text, flags=re.IGNORECASE)
                                    
                                    responsible = re.sub(r'\s+', ' ', resp_text)
                                    responsible = re.sub(r'^:\s*', '', responsible)
                                    collected_resp = True
                                    
                                    for d_keyword in self.date_keywords:
                                        if d_keyword in current:
                                            date_parts = current.split(d_keyword)
                                            if len(date_parts) > 1:
                                                date_text = date_parts[1].strip()
                                                date_text = re.sub(r'^\s*[—–-]?\s*', '', date_text)
                                                due_date_str = re.sub(r'\s+', ' ', date_text)
                                                
                                                date_match = re.search(r'(\d{2}\.\d{2}\.\d{4})', date_text)
                                                if date_match:
                                                    try:
                                                        due_date = datetime.strptime(date_match.group(1), '%d.%m.%Y').date()
                                                    except ValueError:
                                                        pass
                                                collected_date = True
                                            break
                                i += 1
                                break
                    
                    if not collected_date and i < len(task_lines):
                        current = task_lines[i].strip()
                        current_lower = current.lower()
                        
                        # Проверяем на статусы выполнения
                        if any(word in current_lower for word in self.status_keywords):
                            due_date_str = 'Выполнено'
                            collected_date = True
                            i += 1
                        elif 'до конца года' in current_lower:
                            due_date_str = 'до конца года'
                            collected_date = True
                            i += 1
                        else:
                            for keyword in self.date_keywords:
                                if keyword in current:
                                    date_parts = current.split(keyword)
                                    if len(date_parts) > 1:
                                        date_text = date_parts[1].strip()
                                        date_text = re.sub(r'^\s*[—–-]?\s*', '', date_text)
                                        due_date_str = re.sub(r'\s+', ' ', date_text)
                                        
                                        date_match = re.search(r'(\d{2}\.\d{2}\.\d{4})', date_text)
                                        if date_match:
                                            try:
                                                due_date = datetime.strptime(date_match.group(1), '%d.%m.%Y').date()
                                            except ValueError:
                                                pass
                                        collected_date = True
                                        i += 1
                                    break
                    
                    if not (collected_resp or collected_date):
                        i += 1
                
                task = {
                    'number': task_number,
                    'full_description': description,
                    'responsible': responsible,
                    'due_date': due_date,
                    'due_date_str': due_date_str
                }
                tasks.append(task)
                task_counter += 1
            else:
                i += 1
        
        return tasks
    
    def _parse_simple_list(self, lines: List[str]) -> List[Dict]:
        tasks = []
        current_task = None
        current_description = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            task_match = re.match(r'^(\d+)\.\s+(.*)', line)
            
            if task_match:
                if current_task:
                    full_desc = ' '.join(current_description)
                    full_desc = re.sub(r'\s+', ' ', full_desc)
                    current_task['full_description'] = full_desc
                    tasks.append(current_task)
                
                task_num = task_match.group(1)
                task_text = task_match.group(2)
                
                current_task = {
                    'number': int(task_num),
                    'full_description': '',
                    'responsible': '',
                    'due_date': None,
                    'due_date_str': ''
                }
                current_description = [task_text]
            
            elif current_task:
                current_description.append(line)
                
                if 'Срок' in line or any(word in line.lower() for word in self.status_keywords + ['до конца года']):
                    line_lower = line.lower()
                    
                    if any(word in line_lower for word in self.status_keywords):
                        current_task['due_date_str'] = 'Выполнено'
                    elif 'до конца года' in line_lower:
                        current_task['due_date_str'] = 'до конца года'
                    else:
                        date_match = re.search(rf'Срок\s*[—–-]?\s*(\d{{2}}\.\d{{2}}\.\d{{4}})', line)
                        if date_match:
                            date_str = date_match.group(1).strip()
                            current_task['due_date_str'] = date_str
                            try:
                                current_task['due_date'] = datetime.strptime(date_str, '%d.%m.%Y').date()
                            except ValueError:
                                pass
                
                for keyword in self.resp_keywords:
                    if keyword in line:
                        resp_match = re.search(rf'{re.escape(keyword)}\s*[—–-]?\s*([^С]+?)(?:\s+Срок|$)', line)
                        if not resp_match:
                            resp_match = re.search(rf'{re.escape(keyword)}\s*[—–-]?\s*([^\n]+)', line)
                        
                        if resp_match:
                            responsible = resp_match.group(1).strip()
                            stop_words = self.date_keywords + ['Выполнено', 'Приложение', 'приложение'] + self.status_keywords
                            for stop_word in stop_words:
                                if stop_word in responsible.lower():
                                    responsible = responsible.split(stop_word)[0].strip()
                                    break
                            
                            # Дополнительная очистка
                            responsible = re.sub(r'\s+Срок.*$', '', responsible, flags=re.IGNORECASE)
                            responsible = re.sub(r'\s+Выполнено.*$', '', responsible, flags=re.IGNORECASE)
                            responsible = re.sub(r'\s+до\s+конца\s+года.*$', '', responsible, flags=re.IGNORECASE)
                            
                            responsible = re.sub(r'\s+', ' ', responsible)
                            current_task['responsible'] = responsible
                            break
        
        if current_task:
            full_desc = ' '.join(current_description)
            full_desc = re.sub(r'\s+', ' ', full_desc)
            current_task['full_description'] = full_desc
            tasks.append(current_task)
        
        return tasks
    
    def print_tasks(self):
        if not self.tasks:
            print("❌ Задачи не найдены")
            return
        
        print(f"\n📋 Найдено задач: {len(self.tasks)}\n")
        print("=" * 80)
        
        for task in self.tasks:
            print(f"Задача #{task['number']}")
            print(f"📝 Описание: {task['full_description'][:100]}...")
            print(f"👤 Ответственный: {task['responsible'] or '❌ НЕТ'}")
            print(f"📅 Срок: {task['due_date_str'] or '❌ НЕТ'}")
            print("-" * 40)
    
    def to_dataframe(self):
        import pandas as pd
        
        data = []
        for task in self.tasks:
            data.append({
                '№': task['number'],
                'Описание': task['full_description'],
                'Ответственный': task.get('responsible', 'Не указан'),
                'Срок': task.get('due_date_str', 'Не указан'),
                'Дата (для сортировки)': task.get('due_date')
            })
        
        df = pd.DataFrame(data)
        return df